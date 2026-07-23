"""
Claustor AI — Document Parser
Primary: Docling (IBM) — handles tables, reading order, multi-column
Fallback: PyMuPDF — fast, lightweight, text-only
"""

import asyncio
import json
from pathlib import Path

import structlog

logger = structlog.get_logger(__name__)


class ParsedDocument:
    """Structured output from document parsing."""

    def __init__(self):
        self.full_text: str = ""
        self.pages: list[dict] = []
        self.tables: list[dict] = []
        self.metadata: dict = {}
        self.chunks: list[dict] = []
        self.page_count: int = 0
        self.parser_used: str = ""

    def to_dict(self) -> dict:
        return {
            "full_text": self.full_text,
            "pages": self.pages,
            "tables": self.tables,
            "metadata": self.metadata,
            "chunks": self.chunks,
            "page_count": self.page_count,
            "parser_used": self.parser_used,
        }


class DocumentParser:
    """
    Document parser with automatic fallback.

    Priority:
      1. Docling — best quality, handles tables + layout
      2. PyMuPDF — fallback if Docling unavailable on Mac
    """

    def __init__(self):
        self._docling_available = None

    async def _check_docling(self) -> bool:
        """Check if Docling is available."""
        if self._docling_available is None:
            try:
                import docling  # noqa: F401
                self._docling_available = True
                logger.info("docling_available")
            except ImportError:
                self._docling_available = False
                logger.warning(
                    "docling_not_available",
                    fallback="pymupdf",
                    note="Install docling for better table + layout support",
                )
        return self._docling_available

    async def parse(
        self,
        file_bytes: bytes,
        filename: str,
        extract_tables: bool = True,
    ) -> ParsedDocument:
        """
        Parse document and return structured output.

        Args:
            file_bytes: Raw file bytes
            filename: Original filename (for extension detection)
            extract_tables: Whether to extract tables (Starter+ plan)
        """
        ext = Path(filename).suffix.lower()

        if ext in [".pdf"]:
            return await self._parse_pdf(file_bytes, extract_tables)
        elif ext in [".docx", ".doc"]:
            return await self._parse_docx(file_bytes)
        else:
            raise ValueError(f"Unsupported file type: {ext}")

    async def _parse_pdf(
        self,
        file_bytes: bytes,
        extract_tables: bool = True,
    ) -> ParsedDocument:
        """Parse PDF — try Docling first, fall back to PyMuPDF."""

        if await self._check_docling():
            try:
                return await self._parse_with_docling(file_bytes, extract_tables)
            except Exception as e:
                logger.warning("docling_parse_failed", error=str(e), fallback="pymupdf")

        return await self._parse_with_pymupdf(file_bytes)

    async def _parse_with_docling(
        self,
        file_bytes: bytes,
        extract_tables: bool,
    ) -> ParsedDocument:
        """Parse using Docling — handles tables, multi-column, reading order."""
        import io
        from docling.document_converter import DocumentConverter
        from docling.datamodel.base_models import InputFormat
        from docling.datamodel.pipeline_options import PdfPipelineOptions

        loop = asyncio.get_event_loop()

        def _parse():
            options = PdfPipelineOptions(
                do_table_structure=extract_tables,
                do_ocr=False,  # OCR handled separately
            )
            converter = DocumentConverter()
            result = converter.convert(
                io.BytesIO(file_bytes),
                raises_on_error=True,
            )
            return result

        result = await loop.run_in_executor(None, _parse)

        doc = ParsedDocument()
        doc.parser_used = "docling"
        doc.full_text = result.document.export_to_markdown()
        doc.page_count = len(result.document.pages) if result.document.pages else 0

        # Extract tables
        if extract_tables:
            for table in result.document.tables:
                try:
                    df = table.export_to_dataframe()
                    doc.tables.append({
                        "headers": df.columns.tolist(),
                        "rows": df.values.tolist(),
                        "text": table.export_to_markdown(),
                        "page": getattr(table, "page_no", 0),
                    })
                except Exception:
                    pass

        # Extract pages
        if result.document.pages:
            for page in result.document.pages:
                doc.pages.append({
                    "page_number": getattr(page, "page_no", 0),
                    "text": "",  # included in full_text
                })

        # Create chunks for RAG
        doc.chunks = self._create_chunks(doc.full_text, doc.tables)

        logger.info(
            "docling_parse_complete",
            pages=doc.page_count,
            tables=len(doc.tables),
            chunks=len(doc.chunks),
            chars=len(doc.full_text),
        )

        return doc

    async def _parse_with_pymupdf(self, file_bytes: bytes) -> ParsedDocument:
        """Parse using PyMuPDF — fast fallback for Mac development."""
        import io
        import pymupdf  # fitz

        loop = asyncio.get_event_loop()

        def _parse():
            doc_result = ParsedDocument()
            doc_result.parser_used = "pymupdf"

            pdf = pymupdf.open(stream=io.BytesIO(file_bytes), filetype="pdf")
            doc_result.page_count = len(pdf)

            all_text = []
            for page_num, page in enumerate(pdf):
                text = page.get_text("text")
                all_text.append(text)
                doc_result.pages.append({
                    "page_number": page_num + 1,
                    "text": text,
                })

            doc_result.full_text = "\n\n".join(all_text)
            doc_result.chunks = []
            pdf.close()
            return doc_result

        doc = await loop.run_in_executor(None, _parse)
        doc.chunks = self._create_chunks(doc.full_text, [])

        logger.info(
            "pymupdf_parse_complete",
            pages=doc.page_count,
            chunks=len(doc.chunks),
            chars=len(doc.full_text),
        )

        return doc

    async def _parse_docx(self, file_bytes: bytes) -> ParsedDocument:
        """Parse DOCX file."""
        import io
        from docx import Document

        loop = asyncio.get_event_loop()

        def _parse():
            doc_result = ParsedDocument()
            doc_result.parser_used = "python-docx"

            docx = Document(io.BytesIO(file_bytes))
            paragraphs = [p.text for p in docx.paragraphs if p.text.strip()]
            doc_result.full_text = "\n\n".join(paragraphs)

            # Extract tables
            for table in docx.tables:
                rows = []
                for row in table.rows:
                    rows.append([cell.text.strip() for cell in row.cells])
                if rows:
                    doc_result.tables.append({
                        "headers": rows[0] if rows else [],
                        "rows": rows[1:] if len(rows) > 1 else [],
                        "text": "\n".join([" | ".join(row) for row in rows]),
                        "page": 0,
                    })

            return doc_result

        doc = await loop.run_in_executor(None, _parse)
        doc.chunks = self._create_chunks(doc.full_text, doc.tables)

        logger.info(
            "docx_parse_complete",
            tables=len(doc.tables),
            chunks=len(doc.chunks),
        )

        return doc

    def _create_chunks(
        self,
        full_text: str,
        tables: list[dict],
        chunk_size: int = 1000,
        overlap: int = 200,
    ) -> list[dict]:
        """
        Split text into overlapping chunks for RAG indexing.
        Larger chunks = better context but more tokens.
        Overlap = prevents losing context at chunk boundaries.
        """
        chunks = []

        # Text chunks
        text = full_text.strip()
        if not text:
            return chunks

        start = 0
        chunk_index = 0
        while start < len(text):
            end = start + chunk_size

            # Try to break at sentence boundary
            if end < len(text):
                # Look for ". " or "\n" near the end
                for break_char in [". ", "\n\n", "\n", " "]:
                    pos = text.rfind(break_char, start + chunk_size - 100, end)
                    if pos > start:
                        end = pos + len(break_char)
                        break

            chunk_text = text[start:end].strip()
            if chunk_text:
                chunks.append({
                    "text": chunk_text,
                    "chunk_index": chunk_index,
                    "type": "text",
                    "start_char": start,
                    "end_char": end,
                })
                chunk_index += 1

            start = end - overlap  # overlap for context continuity

        # Table chunks — each table is its own chunk
        for table in tables:
            table_text = table.get("text", "")
            if table_text:
                chunks.append({
                    "text": f"[TABLE]\n{table_text}",
                    "chunk_index": chunk_index,
                    "type": "table",
                    "page": table.get("page", 0),
                })
                chunk_index += 1

        return chunks


# Singleton
_parser: DocumentParser | None = None


def get_document_parser() -> DocumentParser:
    """Get or create singleton document parser."""
    global _parser
    if _parser is None:
        _parser = DocumentParser()
    return _parser
